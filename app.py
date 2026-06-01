"""
🖊️ EDITOR PDF by PARBUS
Aplicación independiente - Firma digital de documentos
Servicios Multiples Partner & Business S.A.C.
"""
from flask import Flask, render_template, request, jsonify, send_file, send_from_directory
import os
import uuid
import shutil
from io import BytesIO

app = Flask(__name__)
app.secret_key = "Editor_PDF_Parbus_2026"

BASE_DIR = os.path.dirname(__file__)
UPLOAD_DIR = os.path.join(BASE_DIR, "uploads")
os.makedirs(UPLOAD_DIR, exist_ok=True)

# Colores corporativos
MORADO = "#302b63"
NARANJA = "#e67e22"

# ===== PÁGINA PRINCIPAL =====
@app.route("/")
def index():
    return render_template("editor.html")

# ===== SUBIR PDF =====
@app.route("/api/upload", methods=["POST"])
def api_upload():
    if "pdf" not in request.files:
        return jsonify({"ok": False, "error": "No se envió PDF"}), 400
    
    pdf_file = request.files["pdf"]
    sess_id = str(uuid.uuid4())[:8]
    sess_dir = os.path.join(UPLOAD_DIR, sess_id)
    os.makedirs(sess_dir, exist_ok=True)
    
    pdf_path = os.path.join(sess_dir, "original.pdf")
    pdf_file.save(pdf_path)
    
    try:
        import fitz
        doc = fitz.open(pdf_path)
        paginas = []
        for i in range(doc.page_count):
            page = doc[i]
            pix = page.get_pixmap(matrix=fitz.Matrix(2, 2))
            img_path = os.path.join(sess_dir, f"pag_{i+1}.png")
            pix.save(img_path)
            paginas.append(f"/uploads/{sess_id}/pag_{i+1}.png")
        doc.close()
        return jsonify({"ok": True, "sess_id": sess_id, "paginas": len(paginas)})
    except Exception as e:
        return jsonify({"ok": False, "error": str(e)}), 500

# ===== SUBIR FIRMA =====
@app.route("/api/firma", methods=["POST"])
def api_firma():
    if "firma" not in request.files:
        return jsonify({"ok": False, "error": "No se envió firma"}), 400
    
    f = request.files["firma"]
    firma_dir = os.path.join(BASE_DIR, "static")
    os.makedirs(firma_dir, exist_ok=True)
    firma_path = os.path.join(firma_dir, "firma_activa.png")
    f.save(firma_path)
    return jsonify({"ok": True, "url": "/static/firma_activa.png"})

# ===== OBTENER PÁGINAS COMO IMÁGENES =====
@app.route("/api/<sess_id>/paginas")
def api_paginas(sess_id):
    sess_dir = os.path.join(UPLOAD_DIR, sess_id)
    if not os.path.exists(sess_dir):
        return jsonify({"ok": False, "error": "Sesión no encontrada"}), 404
    
    import glob
    paginas = sorted(glob.glob(os.path.join(sess_dir, "pag_*.png")))
    urls = [f"/uploads/{sess_id}/{os.path.basename(p)}" for p in paginas]
    return jsonify({"ok": True, "paginas": urls})

# ===== FIRMAR PDF =====
@app.route("/api/<sess_id>/firmar", methods=["POST"])
def api_firmar(sess_id):
    data = request.get_json()
    if not data:
        return jsonify({"ok": False, "error": "Datos inválidos"}), 400
    
    sess_dir = os.path.join(UPLOAD_DIR, sess_id)
    pdf_path = os.path.join(sess_dir, "original.pdf")
    firma_path = os.path.join(BASE_DIR, "static", "firma_activa.png")
    
    if not os.path.exists(pdf_path):
        return jsonify({"ok": False, "error": "PDF no encontrado"}), 404
    if not os.path.exists(firma_path):
        return jsonify({"ok": False, "error": "No hay firma cargada. Sube una imagen primero."}), 400
    
    firmas_data = data.get("firmas", {})
    escala = float(data.get("escala", 150))
    
    try:
        import fitz
        doc = fitz.open(pdf_path)
        page_h = doc[0].rect.height
        page_w = doc[0].rect.width
        
        for pag_str, lista in firmas_data.items():
            pagina = int(pag_str) - 1
            if pagina < 0 or pagina >= doc.page_count:
                continue
            
            page = doc[pagina]
            for f in lista:
                x_pct = float(f["x"])
                y_pct = float(f["y"])
                
                sig_w = escala
                sig_h = escala * 0.3
                
                x_pts = (x_pct / 100) * page_w - sig_w / 2
                y_pts = (y_pct / 100) * page_h - sig_h / 2
                
                rect = fitz.Rect(x_pts, y_pts, x_pts + sig_w, y_pts + sig_h)
                page.insert_image(rect, filename=firma_path, overlay=True)
        
        buffer = BytesIO()
        doc.save(buffer, incremental=False, deflate=True)
        doc.close()
        buffer.seek(0)
        
        return send_file(
            buffer,
            mimetype='application/pdf',
            as_attachment=True,
            download_name='DOCUMENTO_FIRMADO.pdf'
        )
    except Exception as e:
        return jsonify({"ok": False, "error": str(e)}), 500

# ===== CONVERTIR: PDF → WORD =====
@app.route("/api/<sess_id>/convertir/word", methods=["POST"])
def api_convertir_word(sess_id):
    sess_dir = os.path.join(UPLOAD_DIR, sess_id)
    pdf_path = os.path.join(sess_dir, "original.pdf")
    if not os.path.exists(pdf_path):
        return jsonify({"ok": False, "error": "PDF no encontrado"}), 404
    try:
        from pdf2docx import Converter
        docx_path = os.path.join(sess_dir, "convertido.docx")
        cv = Converter(pdf_path)
        cv.convert(docx_path, start=0, end=None)
        cv.close()
        return send_file(docx_path, as_attachment=True, download_name="DOCUMENTO.docx")
    except Exception as e:
        return jsonify({"ok": False, "error": str(e)}), 500

# ===== CONVERTIR: PDF → EXCEL (tablas) =====
@app.route("/api/<sess_id>/convertir/excel", methods=["POST"])
def api_convertir_excel(sess_id):
    sess_dir = os.path.join(UPLOAD_DIR, sess_id)
    pdf_path = os.path.join(sess_dir, "original.pdf")
    if not os.path.exists(pdf_path):
        return jsonify({"ok": False, "error": "PDF no encontrado"}), 404
    try:
        import fitz
        from openpyxl import Workbook
        doc = fitz.open(pdf_path)
        wb = Workbook()
        
        for i in range(doc.page_count):
            page = doc[i]
            tabs = page.find_tables()
            if tabs.tables:
                ws = wb.active if i == 0 else wb.create_sheet()
                ws.title = f"Pag_{i+1}"
                for ti, table in enumerate(tabs.tables):
                    data = table.extract()
                    for ri, row in enumerate(data):
                        for ci, cell in enumerate(row):
                            ws.cell(row=ri+1, column=ci+1, value=cell)
        
        doc.close()
        xlsx_path = os.path.join(sess_dir, "convertido.xlsx")
        wb.save(xlsx_path)
        return send_file(xlsx_path, as_attachment=True, download_name="TABLAS_EXTRAIDAS.xlsx")
    except Exception as e:
        return jsonify({"ok": False, "error": str(e)}), 500

# ===== CONVERTIR: PDF → IMÁGENES =====
@app.route("/api/<sess_id>/convertir/imagenes", methods=["POST"])
def api_convertir_imagenes(sess_id):
    sess_dir = os.path.join(UPLOAD_DIR, sess_id)
    pdf_path = os.path.join(sess_dir, "original.pdf")
    if not os.path.exists(pdf_path):
        return jsonify({"ok": False, "error": "PDF no encontrado"}), 404
    try:
        import fitz
        import shutil
        img_dir = os.path.join(sess_dir, "imagenes")
        os.makedirs(img_dir, exist_ok=True)
        doc = fitz.open(pdf_path)
        for i in range(doc.page_count):
            page = doc[i]
            pix = page.get_pixmap(matrix=fitz.Matrix(3, 3))
            img_path = os.path.join(img_dir, f"pag_{i+1}.png")
            pix.save(img_path)
        doc.close()
        # ZIP
        shutil.make_archive(os.path.join(sess_dir, "imagenes"), "zip", img_dir)
        return send_file(os.path.join(sess_dir, "imagenes.zip"), as_attachment=True, download_name="PAGINAS_COMO_IMAGENES.zip")
    except Exception as e:
        return jsonify({"ok": False, "error": str(e)}), 500

# ===== UNIR PDFS =====
@app.route("/api/unir", methods=["POST"])
def api_unir_pdfs():
    files = request.files.getlist("pdfs")
    if len(files) < 2:
        return jsonify({"ok": False, "error": "Sube al menos 2 PDFs"}), 400
    try:
        import fitz
        sess_id = str(uuid.uuid4())[:8]
        sess_dir = os.path.join(UPLOAD_DIR, sess_id)
        os.makedirs(sess_dir, exist_ok=True)
        doc_out = fitz.open()
        for f in files:
            tmp = os.path.join(sess_dir, f"tmp_{f.filename}")
            f.save(tmp)
            doc_in = fitz.open(tmp)
            doc_out.insert_pdf(doc_in)
            doc_in.close()
        merged_path = os.path.join(sess_dir, "unido.pdf")
        doc_out.save(merged_path)
        doc_out.close()
        return send_file(merged_path, as_attachment=True, download_name="PDF_UNIDO.pdf")
    except Exception as e:
        return jsonify({"ok": False, "error": str(e)}), 500

# ===== COMPRIMIR PDF =====
@app.route("/api/<sess_id>/comprimir", methods=["POST"])
def api_comprimir_pdf(sess_id):
    sess_dir = os.path.join(UPLOAD_DIR, sess_id)
    pdf_path = os.path.join(sess_dir, "original.pdf")
    if not os.path.exists(pdf_path):
        return jsonify({"ok": False, "error": "PDF no encontrado"}), 404
    try:
        import fitz
        doc = fitz.open(pdf_path)
        # Guardar con compresión
        comp_path = os.path.join(sess_dir, "comprimido.pdf")
        doc.save(comp_path, garbage=4, deflate=True, clean=True)
        doc.close()
        orig_size = os.path.getsize(pdf_path)
        comp_size = os.path.getsize(comp_path)
        ratio = round((1 - comp_size/orig_size) * 100, 1)
        return send_file(comp_path, as_attachment=True, download_name="PDF_COMPRIMIDO.pdf")
    except Exception as e:
        return jsonify({"ok": False, "error": str(e)}), 500

# ===== DIVIDIR PDF =====
@app.route("/api/<sess_id>/dividir", methods=["POST"])
def api_dividir_pdf(sess_id):
    sess_dir = os.path.join(UPLOAD_DIR, sess_id)
    pdf_path = os.path.join(sess_dir, "original.pdf")
    if not os.path.exists(pdf_path):
        return jsonify({"ok": False, "error": "PDF no encontrado"}), 404
    data = request.get_json() or {}
    paginas = data.get("paginas", "1-")
    try:
        import fitz
        doc = fitz.open(pdf_path)
        out = fitz.open()
        # Parse page ranges like "1-3,5,7-9"
        indices = []
        for part in paginas.split(","):
            part = part.strip()
            if "-" in part:
                a, b = part.split("-")
                for p in range(int(a)-1, int(b)):
                    if 0 <= p < doc.page_count:
                        indices.append(p)
            else:
                p = int(part)-1
                if 0 <= p < doc.page_count:
                    indices.append(p)
        for p in indices:
            out.insert_pdf(doc, from_page=p, to_page=p)
        doc.close()
        out_path = os.path.join(sess_dir, "dividido.pdf")
        out.save(out_path)
        out.close()
        return send_file(out_path, as_attachment=True, download_name="PDF_DIVIDIDO.pdf")
    except Exception as e:
        return jsonify({"ok": False, "error": str(e)}), 500

# ===== JPG → PDF =====
@app.route("/api/imagenes-a-pdf", methods=["POST"])
def api_imagenes_pdf():
    files = request.files.getlist("imagenes")
    if not files:
        return jsonify({"ok": False, "error": "Sube al menos una imagen"}), 400
    try:
        import fitz
        sess_id = str(uuid.uuid4())[:8]
        sess_dir = os.path.join(UPLOAD_DIR, sess_id)
        os.makedirs(sess_dir, exist_ok=True)
        doc = fitz.open()
        for f in sorted(files, key=lambda x: x.filename):
            img_bytes = f.read()
            img_pdf = fitz.open(stream=img_bytes, filetype="png" if f.filename.lower().endswith(".png") else "jpeg")
            doc.insert_pdf(img_pdf)
            img_pdf.close()
        out_path = os.path.join(sess_dir, "imagenes.pdf")
        doc.save(out_path)
        doc.close()
        return send_file(out_path, as_attachment=True, download_name="IMAGENES.pdf")
    except Exception as e:
        return jsonify({"ok": False, "error": str(e)}), 500

# ===== ROTAR PDF =====
@app.route("/api/<sess_id>/rotar", methods=["POST"])
def api_rotar_pdf(sess_id):
    sess_dir = os.path.join(UPLOAD_DIR, sess_id)
    pdf_path = os.path.join(sess_dir, "original.pdf")
    if not os.path.exists(pdf_path):
        return jsonify({"ok": False, "error": "PDF no encontrado"}), 404
    data = request.get_json() or {}
    angulo = int(data.get("angulo", 90))
    paginas = data.get("paginas", "todas")
    try:
        import fitz
        doc = fitz.open(pdf_path)
        if paginas == "todas":
            for page in doc:
                page.set_rotation(page.rotation + angulo)
        else:
            for p in paginas:
                if 0 <= p-1 < doc.page_count:
                    doc[p-1].set_rotation(doc[p-1].rotation + angulo)
        out_path = os.path.join(sess_dir, "rotado.pdf")
        doc.save(out_path)
        doc.close()
        return send_file(out_path, as_attachment=True, download_name="PDF_ROTADO.pdf")
    except Exception as e:
        return jsonify({"ok": False, "error": str(e)}), 500

# ===== ELIMINAR PÁGINAS =====
@app.route("/api/<sess_id>/eliminar-paginas", methods=["POST"])
def api_eliminar_paginas(sess_id):
    sess_dir = os.path.join(UPLOAD_DIR, sess_id)
    pdf_path = os.path.join(sess_dir, "original.pdf")
    if not os.path.exists(pdf_path):
        return jsonify({"ok": False, "error": "PDF no encontrado"}), 404
    data = request.get_json() or {}
    eliminar = data.get("paginas", [])
    try:
        import fitz
        doc = fitz.open(pdf_path)
        eliminar_set = set(int(p)-1 for p in eliminar if 1 <= int(p) <= doc.page_count)
        out = fitz.open()
        for i in range(doc.page_count):
            if i not in eliminar_set:
                out.insert_pdf(doc, from_page=i, to_page=i)
        doc.close()
        out_path = os.path.join(sess_dir, "reducido.pdf")
        out.save(out_path)
        out.close()
        return send_file(out_path, as_attachment=True, download_name="PDF_SIN_PAGINAS.pdf")
    except Exception as e:
        return jsonify({"ok": False, "error": str(e)}), 500

# ===== WORD → PDF =====
@app.route("/api/word-a-pdf", methods=["POST"])
def api_word_a_pdf():
    if "doc" not in request.files:
        return jsonify({"ok": False, "error": "No se envió archivo"}), 400
    f = request.files["doc"]
    sess_id = str(uuid.uuid4())[:8]
    sess_dir = os.path.join(UPLOAD_DIR, sess_id)
    os.makedirs(sess_dir, exist_ok=True)
    docx_path = os.path.join(sess_dir, "input.docx")
    f.save(docx_path)
    try:
        # Convertir usando python-docx + fitz (texto básico)
        from docx import Document
        import fitz
        doc = Document(docx_path)
        pdf_doc = fitz.open()
        full_text = []
        for p in doc.paragraphs:
            if p.text.strip():
                full_text.append(p.text)
        # Crear páginas (aprox 50 líneas por página)
        lineas_por_pag = 45
        for i in range(0, len(full_text), lineas_por_pag):
            page = pdf_doc.new_page()
            bloque = "\n".join(full_text[i:i+lineas_por_pag])
            rect = fitz.Rect(50, 50, page.rect.width-50, page.rect.height-50)
            page.insert_text(rect, bloque, fontsize=11)
        out_path = os.path.join(sess_dir, "convertido.pdf")
        pdf_doc.save(out_path)
        pdf_doc.close()
        return send_file(out_path, as_attachment=True, download_name="DOCUMENTO.pdf")
    except Exception as e:
        return jsonify({"ok": False, "error": str(e)}), 500

# ===== EXCEL → PDF =====
@app.route("/api/excel-a-pdf", methods=["POST"])
def api_excel_a_pdf():
    if "xlsx" not in request.files:
        return jsonify({"ok": False, "error": "No se envió archivo"}), 400
    f = request.files["xlsx"]
    sess_id = str(uuid.uuid4())[:8]
    sess_dir = os.path.join(UPLOAD_DIR, sess_id)
    os.makedirs(sess_dir, exist_ok=True)
    xlsx_path = os.path.join(sess_dir, "input.xlsx")
    f.save(xlsx_path)
    try:
        from openpyxl import load_workbook
        import fitz
        wb = load_workbook(xlsx_path, data_only=True)
        pdf_doc = fitz.open()
        for ws_name in wb.sheetnames:
            ws = wb[ws_name]
            page = pdf_doc.new_page()
            y = 50
            page.insert_text(fitz.Rect(50, y, 500, y+15), f"Hoja: {ws_name}", fontsize=14, bold=True)
            y += 25
            for row in ws.iter_rows(values_only=True):
                vals = [str(v)[:20] if v is not None else "" for v in row[:8]]
                linea = " | ".join(vals)
                page.insert_text(fitz.Rect(50, y, 500, y+12), linea, fontsize=8)
                y += 12
                if y > 800:
                    page = pdf_doc.new_page()
                    y = 50
        out_path = os.path.join(sess_dir, "convertido.pdf")
        pdf_doc.save(out_path)
        pdf_doc.close()
        return send_file(out_path, as_attachment=True, download_name="EXCEL_CONVERTIDO.pdf")
    except Exception as e:
        return jsonify({"ok": False, "error": str(e)}), 500

# ===== ARCHIVOS ESTÁTICOS =====
@app.route("/uploads/<path:filename>")
def serve_uploads(filename):
    return send_from_directory(UPLOAD_DIR, filename)

if __name__ == "__main__":
    port = int(os.environ.get("PORT", 5001))
    print("🖊️  EDITOR PDF by PARBUS")
    print(f"   http://localhost:{port}")
    print("   ─────────────────────")
    app.run(debug=False, host="0.0.0.0", port=port)
